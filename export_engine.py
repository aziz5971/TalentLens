"""
export_engine.py  –  Multi-Format Report Export Engine.

Generates professional reports in multiple formats:
  - PDF   (via WeasyPrint HTML→PDF or reportlab fallback)
  - DOCX  (via python-docx)
  - CSV   (for spreadsheet import)
  - Markdown
  - JSON  (already exists in report_generator.py; this adds structured export)

Report types:
  - Single Candidate Report  — full detail for one candidate
  - Comparative Report        — side-by-side ranking of all candidates
  - Executive Summary         — 1-page overview for hiring manager
  - Pipeline Report           — current pipeline state + funnel metrics
"""
from __future__ import annotations

import csv
import io
import json
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Optional

from models import CandidateScore, JDCriteria


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def export_candidate_report(
    score: CandidateScore,
    jd: JDCriteria,
    fmt: str = "md",
    output_dir: Path | None = None,
) -> str:
    """Export a single candidate report. Returns content string or file path."""
    output_dir = output_dir or Path(__file__).parent / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = _safe_filename(score.candidate.name)

    if fmt == "md":
        content = _candidate_markdown(score, jd)
        path = output_dir / f"report_{safe_name}_{ts}.md"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "json":
        content = _candidate_json(score, jd)
        path = output_dir / f"report_{safe_name}_{ts}.json"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "docx":
        path = output_dir / f"report_{safe_name}_{ts}.docx"
        _candidate_docx(score, jd, path)
        return str(path)
    elif fmt == "html":
        content = _candidate_html(score, jd)
        path = output_dir / f"report_{safe_name}_{ts}.html"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "pdf":
        path = output_dir / f"report_{safe_name}_{ts}.pdf"
        _candidate_pdf(score, jd, path)
        return str(path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")


def export_comparative_report(
    scores: list[CandidateScore],
    jd: JDCriteria,
    fmt: str = "md",
    output_dir: Path | None = None,
) -> str:
    """Export comparative ranking report."""
    output_dir = output_dir or Path(__file__).parent / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = _safe_filename(jd.title)

    if fmt == "md":
        content = _comparative_markdown(scores, jd)
        path = output_dir / f"comparative_{safe_title}_{ts}.md"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "csv":
        content = _comparative_csv(scores, jd)
        path = output_dir / f"comparative_{safe_title}_{ts}.csv"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "html":
        content = _comparative_html(scores, jd)
        path = output_dir / f"comparative_{safe_title}_{ts}.html"
        path.write_text(content, encoding="utf-8")
        return str(path)
    elif fmt == "json":
        data = {
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "job_title": jd.title,
            "total_candidates": len(scores),
            "candidates": [_score_to_summary(s) for s in scores],
        }
        content = json.dumps(data, indent=2, default=str)
        path = output_dir / f"comparative_{safe_title}_{ts}.json"
        path.write_text(content, encoding="utf-8")
        return str(path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")


def export_executive_summary(
    scores: list[CandidateScore],
    jd: JDCriteria,
    output_dir: Path | None = None,
) -> str:
    """Generate 1-page executive summary in Markdown."""
    output_dir = output_dir or Path(__file__).parent / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    content = _executive_summary_md(scores, jd)
    path = output_dir / f"executive_summary_{ts}.md"
    path.write_text(content, encoding="utf-8")
    return str(path)


# ---------------------------------------------------------------------------
# Markdown generators
# ---------------------------------------------------------------------------

def _candidate_markdown(s: CandidateScore, jd: JDCriteria) -> str:
    """Full single-candidate report in Markdown."""
    b = s.breakdown
    lines = [
        f"# Candidate Report: {s.candidate.name}",
        f"",
        f"**Position:** {jd.title}",
        f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"**Overall Score:** {s.overall_score}/100 ({s.grade})",
        f"**Trust Score:** {s.verification.overall_trust_score:.0%}",
        f"",
        f"---",
        f"",
        f"## Score Breakdown",
        f"",
        f"| Component | Score | Max |",
        f"|-----------|------:|----:|",
        f"| Required Skills | {b.required_skills:.1f} | 35 |",
        f"| Preferred Skills | {b.preferred_skills:.1f} | 15 |",
        f"| Experience | {b.experience:.1f} | 25 |",
        f"| Education | {b.education:.1f} | 10 |",
        f"| Certifications | {b.certifications:.1f} | 10 |",
        f"| Semantic Match | {b.semantic_similarity:.1f} | 5 |",
        f"| **Total** | **{s.overall_score:.1f}** | **100** |",
        f"",
        f"## Candidate Profile",
        f"",
        f"- **Name:** {s.candidate.name}",
        f"- **Email:** {s.candidate.email}",
        f"- **Phone:** {s.candidate.phone}",
        f"- **Location:** {s.candidate.location}",
        f"- **LinkedIn:** {s.candidate.linkedin_url or 'Not provided'}",
        f"- **GitHub:** {s.candidate.github_url or 'Not provided'}",
        f"- **Total Experience:** {s.candidate.total_experience_years:.1f} years",
        f"",
        f"## Skills Analysis",
        f"",
    ]

    if s.matched_required_skills:
        lines.append(f"### Matched Required Skills ({len(s.matched_required_skills)})")
        for skill in s.matched_required_skills:
            lines.append(f"- ✅ {skill}")
        lines.append("")

    if s.missing_required_skills:
        lines.append(f"### Missing Required Skills ({len(s.missing_required_skills)})")
        for skill in s.missing_required_skills:
            lines.append(f"- ❌ {skill}")
        lines.append("")

    if s.matched_preferred_skills:
        lines.append(f"### Matched Preferred Skills ({len(s.matched_preferred_skills)})")
        for skill in s.matched_preferred_skills:
            lines.append(f"- ✅ {skill}")
        lines.append("")

    # Experience
    lines.extend([
        f"## Experience ({len(s.candidate.experience)} roles)",
        "",
    ])
    for exp in s.candidate.experience:
        lines.append(f"### {exp.title} — {exp.company}")
        lines.append(f"*{exp.start_date} – {exp.end_date} ({exp.duration_months} months)*")
        if exp.location:
            lines.append(f"📍 {exp.location}")
        if exp.description:
            lines.append(f"\n{exp.description}")
        lines.append("")

    # Verification summary
    lines.extend([
        "## Verification Summary",
        "",
        f"**Overall Trust Score:** {s.verification.overall_trust_score:.0%}",
        "",
    ])

    # LinkedIn
    li = s.verification.linkedin
    if li:
        status = "✅ Verified" if li.url_resolves else "❌ Not found"
        lines.append(f"- **LinkedIn:** {status} (authenticity: {li.authenticity_score:.0%})")
        for flag in (li.red_flags or []):
            lines.append(f"  - ⚠️ {flag}")

    # Companies
    if s.verification.companies:
        lines.append("")
        lines.append("### Company Verification")
        lines.append("")
        lines.append("| Company | Verified | Type | Employees | Score |")
        lines.append("|---------|----------|------|-----------|------:|")
        for co in s.verification.companies:
            mark = "✅" if co.found else "❌"
            lines.append(f"| {co.name} | {mark} | {co.company_type} | {co.employee_count} | {co.legitimacy_score:.0%} |")

    lines.append("")
    lines.append("---")
    lines.append(f"*Generated by TalentLens v3.0 on {datetime.now().strftime('%Y-%m-%d %H:%M')}*")

    return "\n".join(lines)


def _candidate_json(s: CandidateScore, jd: JDCriteria) -> str:
    """Structured JSON export for a single candidate."""
    data = {
        "report_type": "single_candidate",
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "position": jd.title,
        "candidate": {
            "name": s.candidate.name,
            "email": s.candidate.email,
            "phone": s.candidate.phone,
            "location": s.candidate.location,
            "linkedin": s.candidate.linkedin_url,
            "github": s.candidate.github_url,
            "experience_years": s.candidate.total_experience_years,
            "skills": s.candidate.skills,
        },
        "score": {
            "overall": s.overall_score,
            "grade": s.grade,
            "breakdown": {
                "required_skills": s.breakdown.required_skills,
                "preferred_skills": s.breakdown.preferred_skills,
                "experience": s.breakdown.experience,
                "education": s.breakdown.education,
                "certifications": s.breakdown.certifications,
                "semantic": s.breakdown.semantic_similarity,
            },
        },
        "skills_analysis": {
            "matched_required": s.matched_required_skills,
            "missing_required": s.missing_required_skills,
            "matched_preferred": s.matched_preferred_skills,
        },
        "verification": {
            "trust_score": s.verification.overall_trust_score,
        },
    }
    return json.dumps(data, indent=2, default=str)


# ---------------------------------------------------------------------------
# Comparative reports
# ---------------------------------------------------------------------------

def _comparative_markdown(scores: list[CandidateScore], jd: JDCriteria) -> str:
    """Side-by-side ranking Markdown report."""
    lines = [
        f"# Comparative Candidate Report",
        f"",
        f"**Position:** {jd.title}",
        f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"**Candidates Evaluated:** {len(scores)}",
        f"",
        f"---",
        f"",
        f"## Rankings",
        f"",
        f"| # | Candidate | Score | Grade | Req Skills | Experience | Trust |",
        f"|:-:|-----------|------:|:-----:|-----------:|-----------:|------:|",
    ]

    for s in scores:
        req_total = len(s.matched_required_skills) + len(s.missing_required_skills)
        req_pct = len(s.matched_required_skills) / max(req_total, 1) * 100
        lines.append(
            f"| {s.rank} | {s.candidate.name} | {s.overall_score:.1f} | {s.grade} "
            f"| {req_pct:.0f}% ({len(s.matched_required_skills)}/{req_total}) "
            f"| {s.candidate.total_experience_years:.1f} yrs "
            f"| {s.verification.overall_trust_score:.0%} |"
        )

    # Top candidate detail
    if scores:
        top = scores[0]
        lines.extend([
            "",
            "---",
            "",
            f"## Top Candidate: {top.candidate.name}",
            "",
            f"**Score:** {top.overall_score}/100 ({top.grade})",
            f"**Experience:** {top.candidate.total_experience_years:.1f} years",
            f"**Trust:** {top.verification.overall_trust_score:.0%}",
            "",
        ])
        if top.matched_required_skills:
            lines.append(f"**Matched Skills:** {', '.join(top.matched_required_skills)}")
        if top.missing_required_skills:
            lines.append(f"**Missing Skills:** {', '.join(top.missing_required_skills)}")

    lines.extend([
        "",
        "---",
        f"*Generated by TalentLens v3.0 on {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
    ])

    return "\n".join(lines)


def _comparative_csv(scores: list[CandidateScore], jd: JDCriteria) -> str:
    """CSV export for spreadsheet import."""
    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow([
        "Rank", "Name", "Email", "Score", "Grade",
        "Required Skills Match %", "Experience (yrs)",
        "Trust Score", "LinkedIn Verified",
        "Matched Required Skills", "Missing Required Skills",
        "Source File",
    ])

    for s in scores:
        req_total = len(s.matched_required_skills) + len(s.missing_required_skills)
        req_pct = len(s.matched_required_skills) / max(req_total, 1) * 100

        li_verified = "Yes" if (s.verification.linkedin and s.verification.linkedin.url_resolves) else "No"

        writer.writerow([
            s.rank,
            s.candidate.name,
            s.candidate.email,
            f"{s.overall_score:.1f}",
            s.grade,
            f"{req_pct:.0f}%",
            f"{s.candidate.total_experience_years:.1f}",
            f"{s.verification.overall_trust_score:.0%}",
            li_verified,
            "; ".join(s.matched_required_skills),
            "; ".join(s.missing_required_skills),
            s.candidate.source_file,
        ])

    return output.getvalue()


def _comparative_html(scores: list[CandidateScore], jd: JDCriteria) -> str:
    """HTML comparative report with professional styling."""
    rows = ""
    for s in scores:
        req_total = len(s.matched_required_skills) + len(s.missing_required_skills)
        req_pct = len(s.matched_required_skills) / max(req_total, 1) * 100
        grade_color = _grade_color(s.grade)
        trust = s.verification.overall_trust_score

        rows += f"""
        <tr>
            <td class="rank">{s.rank}</td>
            <td class="name">{_html_escape(s.candidate.name)}</td>
            <td class="score"><span class="grade" style="background:{grade_color}">{s.grade}</span> {s.overall_score:.1f}</td>
            <td>{req_pct:.0f}%</td>
            <td>{s.candidate.total_experience_years:.1f} yrs</td>
            <td class="trust">{trust:.0%}</td>
        </tr>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>TalentLens — {_html_escape(jd.title)}</title>
<style>
  :root {{
    --bg: #0f172a; --surface: #1e293b; --border: #334155;
    --text: #f1f5f9; --text-dim: #94a3b8; --accent: #38bdf8;
    --green: #22c55e; --yellow: #eab308; --red: #ef4444;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    background: var(--bg); color: var(--text);
    padding: 2rem; line-height: 1.6;
  }}
  .header {{
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid var(--border); border-radius: 12px;
    padding: 2rem; margin-bottom: 2rem;
  }}
  .header h1 {{ font-size: 1.8rem; font-weight: 700; margin-bottom: 0.5rem; }}
  .header .meta {{ color: var(--text-dim); font-size: 0.9rem; }}
  table {{
    width: 100%; border-collapse: collapse;
    background: var(--surface); border-radius: 12px;
    overflow: hidden; border: 1px solid var(--border);
  }}
  th {{
    background: rgba(56, 189, 248, 0.1); color: var(--accent);
    padding: 1rem; text-align: left; font-weight: 600;
    font-size: 0.85rem; text-transform: uppercase; letter-spacing: 0.05em;
  }}
  td {{ padding: 0.85rem 1rem; border-top: 1px solid var(--border); }}
  tr:hover {{ background: rgba(56, 189, 248, 0.05); }}
  .rank {{ font-weight: 700; color: var(--accent); text-align: center; width: 3rem; }}
  .name {{ font-weight: 600; }}
  .score {{ font-weight: 600; }}
  .grade {{
    display: inline-block; padding: 0.15rem 0.5rem;
    border-radius: 4px; font-size: 0.8rem; font-weight: 700;
    color: white; margin-right: 0.5rem;
  }}
  .trust {{ font-weight: 600; }}
  .footer {{
    margin-top: 2rem; text-align: center;
    color: var(--text-dim); font-size: 0.8rem;
  }}
</style>
</head>
<body>
<div class="header">
    <h1>Candidate Comparison Report</h1>
    <div class="meta">
        Position: {_html_escape(jd.title)} &nbsp;|&nbsp;
        Candidates: {len(scores)} &nbsp;|&nbsp;
        Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}
    </div>
</div>
<table>
    <thead>
        <tr>
            <th>#</th><th>Candidate</th><th>Score</th>
            <th>Skills Match</th><th>Experience</th><th>Trust</th>
        </tr>
    </thead>
    <tbody>
        {rows}
    </tbody>
</table>
<div class="footer">Generated by TalentLens v3.0</div>
</body>
</html>"""


def _candidate_html(s: CandidateScore, jd: JDCriteria) -> str:
    """Single candidate HTML report."""
    b = s.breakdown
    grade_color = _grade_color(s.grade)

    matched_html = "".join(f'<span class="skill matched">{_html_escape(sk)}</span>' for sk in s.matched_required_skills)
    missing_html = "".join(f'<span class="skill missing">{_html_escape(sk)}</span>' for sk in s.missing_required_skills)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>TalentLens — {_html_escape(s.candidate.name)}</title>
<style>
  :root {{
    --bg: #0f172a; --surface: #1e293b; --border: #334155;
    --text: #f1f5f9; --text-dim: #94a3b8; --accent: #38bdf8;
    --green: #22c55e; --red: #ef4444;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: 'Inter', -apple-system, sans-serif;
    background: var(--bg); color: var(--text);
    padding: 2rem; line-height: 1.6; max-width: 900px; margin: 0 auto;
  }}
  .card {{
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 12px; padding: 1.5rem; margin-bottom: 1.5rem;
  }}
  .card h2 {{ color: var(--accent); margin-bottom: 1rem; font-size: 1.2rem; }}
  .hero {{ text-align: center; padding: 2.5rem; }}
  .hero h1 {{ font-size: 2rem; margin-bottom: 0.5rem; }}
  .hero .score {{
    font-size: 3rem; font-weight: 800; color: {grade_color};
  }}
  .hero .grade {{
    display: inline-block; padding: 0.25rem 1rem;
    background: {grade_color}; color: white;
    border-radius: 8px; font-weight: 700; font-size: 1.2rem;
  }}
  .bar {{ height: 8px; border-radius: 4px; background: var(--border); margin: 0.3rem 0; }}
  .bar-fill {{ height: 100%; border-radius: 4px; background: var(--accent); }}
  .breakdown-row {{ display: flex; justify-content: space-between; align-items: center; margin: 0.5rem 0; }}
  .breakdown-label {{ min-width: 140px; color: var(--text-dim); }}
  .breakdown-bar {{ flex: 1; margin: 0 1rem; }}
  .breakdown-val {{ min-width: 60px; text-align: right; font-weight: 600; }}
  .skill {{
    display: inline-block; padding: 0.2rem 0.6rem; margin: 0.2rem;
    border-radius: 6px; font-size: 0.85rem; font-weight: 500;
  }}
  .skill.matched {{ background: rgba(34,197,94,0.15); color: var(--green); border: 1px solid rgba(34,197,94,0.3); }}
  .skill.missing {{ background: rgba(239,68,68,0.15); color: var(--red); border: 1px solid rgba(239,68,68,0.3); }}
  .footer {{ text-align: center; color: var(--text-dim); font-size: 0.8rem; margin-top: 2rem; }}
</style>
</head>
<body>
<div class="card hero">
    <h1>{_html_escape(s.candidate.name)}</h1>
    <p style="color:var(--text-dim)">{_html_escape(jd.title)}</p>
    <div class="score">{s.overall_score:.1f}</div>
    <div class="grade">{s.grade}</div>
    <p style="color:var(--text-dim);margin-top:0.5rem">Trust: {s.verification.overall_trust_score:.0%}</p>
</div>

<div class="card">
    <h2>Score Breakdown</h2>
    {"".join(_score_bar_html(label, val, mx) for label, val, mx in [
        ("Required Skills", b.required_skills, 35),
        ("Preferred Skills", b.preferred_skills, 15),
        ("Experience", b.experience, 25),
        ("Education", b.education, 10),
        ("Certifications", b.certifications, 10),
        ("Semantic Match", b.semantic_similarity, 5),
    ])}
</div>

<div class="card">
    <h2>Skills Analysis</h2>
    <h3 style="color:var(--green);margin-bottom:0.5rem">Matched Required ({len(s.matched_required_skills)})</h3>
    <div>{matched_html or '<span style="color:var(--text-dim)">None</span>'}</div>
    <h3 style="color:var(--red);margin:1rem 0 0.5rem">Missing Required ({len(s.missing_required_skills)})</h3>
    <div>{missing_html or '<span style="color:var(--text-dim)">None</span>'}</div>
</div>

<div class="footer">Generated by TalentLens v3.0 — {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
</body>
</html>"""


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

def _candidate_docx(s: CandidateScore, jd: JDCriteria, path: Path) -> None:
    """Generate a DOCX report for a single candidate."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback: write markdown instead
        md = _candidate_markdown(s, jd)
        path.with_suffix(".md").write_text(md, encoding="utf-8")
        return

    doc = Document()

    # Title
    title = doc.add_heading(f"Candidate Report: {s.candidate.name}", level=1)
    doc.add_paragraph(f"Position: {jd.title}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Overall Score: {s.overall_score}/100 ({s.grade})")
    doc.add_paragraph(f"Trust Score: {s.verification.overall_trust_score:.0%}")

    # Score breakdown table
    doc.add_heading("Score Breakdown", level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = "Light List Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "Component"
    headers[1].text = "Score"
    headers[2].text = "Max"

    b = s.breakdown
    data = [
        ("Required Skills", b.required_skills, 35),
        ("Preferred Skills", b.preferred_skills, 15),
        ("Experience", b.experience, 25),
        ("Education", b.education, 10),
        ("Certifications", b.certifications, 10),
        ("Semantic Match", b.semantic_similarity, 5),
    ]
    for i, (label, val, mx) in enumerate(data):
        row = table.rows[i + 1].cells
        row[0].text = label
        row[1].text = f"{val:.1f}"
        row[2].text = str(mx)

    # Skills
    doc.add_heading("Skills Analysis", level=2)
    if s.matched_required_skills:
        doc.add_paragraph(f"Matched Required: {', '.join(s.matched_required_skills)}")
    if s.missing_required_skills:
        doc.add_paragraph(f"Missing Required: {', '.join(s.missing_required_skills)}")

    # Experience
    doc.add_heading("Experience", level=2)
    for exp in s.candidate.experience:
        doc.add_heading(f"{exp.title} — {exp.company}", level=3)
        doc.add_paragraph(f"{exp.start_date} – {exp.end_date} ({exp.duration_months} months)")
        if exp.description:
            doc.add_paragraph(exp.description)

    doc.save(str(path))


# ---------------------------------------------------------------------------
# PDF generator
# ---------------------------------------------------------------------------

def _candidate_pdf(s: CandidateScore, jd: JDCriteria, path: Path) -> None:
    """Generate PDF from HTML using WeasyPrint or fallback."""
    html_content = _candidate_html(s, jd)

    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(str(path))
    except ImportError:
        # Fallback: save HTML (user can print to PDF from browser)
        html_path = path.with_suffix(".html")
        html_path.write_text(html_content, encoding="utf-8")


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def _executive_summary_md(scores: list[CandidateScore], jd: JDCriteria) -> str:
    """1-page executive summary for hiring managers."""
    total = len(scores)
    avg_score = sum(s.overall_score for s in scores) / max(total, 1)
    top_3 = scores[:3]
    grade_dist = {}
    for s in scores:
        grade_dist[s.grade] = grade_dist.get(s.grade, 0) + 1

    lines = [
        f"# Executive Summary — {jd.title}",
        f"",
        f"**Date:** {datetime.now().strftime('%Y-%m-%d')}",
        f"**Candidates Screened:** {total}",
        f"**Average Score:** {avg_score:.1f}/100",
        f"**Required Skills:** {len(jd.required_skills)}",
        f"**Experience Required:** {jd.min_experience_years}+ years",
        f"",
        f"## Grade Distribution",
        f"",
    ]

    for grade in ["A+", "A", "B+", "B", "C", "D", "F"]:
        count = grade_dist.get(grade, 0)
        if count > 0:
            bar = "█" * count
            lines.append(f"  {grade:3s} {bar} ({count})")

    lines.extend([
        f"",
        f"## Top Candidates",
        f"",
    ])

    for i, s in enumerate(top_3, 1):
        lines.extend([
            f"### {i}. {s.candidate.name} — {s.overall_score:.1f}/100 ({s.grade})",
            f"- Experience: {s.candidate.total_experience_years:.1f} years",
            f"- Skills Match: {len(s.matched_required_skills)}/{len(s.matched_required_skills) + len(s.missing_required_skills)} required",
            f"- Trust: {s.verification.overall_trust_score:.0%}",
            f"",
        ])

    lines.extend([
        f"## Recommendation",
        f"",
    ])

    if top_3 and top_3[0].overall_score >= 80:
        lines.append(f"**Strong candidate pool.** {top_3[0].candidate.name} is a clear frontrunner. Recommend proceeding to interviews immediately.")
    elif top_3 and top_3[0].overall_score >= 60:
        lines.append(f"**Moderate match quality.** Consider interviewing top 2-3 candidates while continuing to source.")
    else:
        lines.append(f"**Weak candidate pool.** Recommend revising JD requirements or expanding search channels.")

    lines.extend([
        f"",
        f"---",
        f"*Generated by TalentLens v3.0*",
    ])

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe_filename(name: str) -> str:
    """Sanitize a name for use in filenames."""
    return re.sub(r"[^\w\-]", "_", name)[:50].strip("_")


import re


def _html_escape(text: str) -> str:
    """Basic HTML escaping."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#x27;")
    )


def _grade_color(grade: str) -> str:
    return {
        "A+": "#22c55e", "A": "#22c55e",
        "B+": "#38bdf8", "B": "#38bdf8",
        "C": "#eab308", "D": "#f97316", "F": "#ef4444",
    }.get(grade, "#94a3b8")


def _score_bar_html(label: str, value: float, max_val: int) -> str:
    pct = min(100, value / max(max_val, 1) * 100)
    return f"""
    <div class="breakdown-row">
        <span class="breakdown-label">{label}</span>
        <div class="breakdown-bar"><div class="bar"><div class="bar-fill" style="width:{pct:.0f}%"></div></div></div>
        <span class="breakdown-val">{value:.1f}/{max_val}</span>
    </div>"""


def _score_to_summary(s: CandidateScore) -> dict:
    return {
        "rank": s.rank,
        "name": s.candidate.name,
        "score": s.overall_score,
        "grade": s.grade,
        "experience_years": s.candidate.total_experience_years,
        "trust_score": s.verification.overall_trust_score,
    }
